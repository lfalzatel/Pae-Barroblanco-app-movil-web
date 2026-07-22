import os
from docx import Document

def update_docx(file_path):
    doc = Document(file_path)
    replacements = {
        "10 de noviembre de 2025": "5 de diciembre de 2025",
        "10 de mayo de 2026": "31 de mayo de 2026",
        "Henry Alirio Mejía Zuluaga": "Cathalina Ceballos Toro"
    }
    
    changed = False
    
    # Check paragraphs
    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                p.text = p.text.replace(old, new)
                changed = True
                
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in p.text:
                            p.text = p.text.replace(old, new)
                            changed = True

    if changed:
        out_path = file_path.replace(".docx", "_Actualizado.docx")
        doc.save(out_path)
        print("Documento actualizado guardado como:", out_path)
    else:
        print("No se encontraron coincidencias para actualizar.")

if __name__ == "__main__":
    file = r"1. Documentos practica\Plan_Practica_SMEM_Luis_Alzate.docx"
    update_docx(file)
